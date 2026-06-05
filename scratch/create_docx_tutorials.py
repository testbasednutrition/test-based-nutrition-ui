import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    # Helper to set background color of a table cell
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def apply_text_formatting(run, font_name="Arial", size_pt=10, color_rgb=None, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_heading_styled(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    
    if level == 1:
        apply_text_formatting(run, font_name="Arial", size_pt=18, color_rgb=RGBColor(139, 26, 26), bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        apply_text_formatting(run, font_name="Arial", size_pt=14, color_rgb=RGBColor(139, 26, 26), bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        apply_text_formatting(run, font_name="Arial", size_pt=11.5, color_rgb=RGBColor(50, 50, 50), bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p

def add_bullet_point(doc, bold_prefix, text_body):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.4)
    
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        apply_text_formatting(run_bold, font_name="Arial", size_pt=10, bold=True)
    
    run_text = p.add_run(text_body)
    apply_text_formatting(run_text, font_name="Arial", size_pt=10)
    return p

def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F7F5F0") # Sand/cream tint background
    
    # Left border formatting is tricky in python-docx, so we'll rely on text styling and padding
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(text)
    apply_text_formatting(run, font_name="Arial", size_pt=9.5, color_rgb=RGBColor(139, 26, 26), italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(4) # Spacing after table

def build_admin_tutorial():
    doc = Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Page / Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("TESTBASED NUTRITION (TBN)\n")
    apply_text_formatting(run_title, font_name="Arial", size_pt=24, color_rgb=RGBColor(139, 26, 26), bold=True)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    run_sub = p_sub.add_run("System Administrator & Back-Office Operations Guide")
    apply_text_formatting(run_sub, font_name="Arial", size_pt=14, color_rgb=RGBColor(100, 100, 100), italic=True)

    add_heading_styled(doc, "Ecosystem Overview", 1)
    p_intro = doc.add_paragraph()
    p_intro_run = p_intro.add_run(
        "Welcome to the TestBased Nutrition (TBN) Administrator Portal Manual. "
        "As an administrator, you oversee the clinical data flow, verify specialists, "
        "manage the CRM client pipelines, track self-hosted web traffic, and distribute marketing "
        "assets to partners. This guide walks you through each section of the B2B administrative portal "
        "and B2C directories."
    )
    apply_text_formatting(p_intro_run, font_name="Arial", size_pt=10)

    # Section 1
    add_heading_styled(doc, "1. Specialist Directory & Approvals (Manage Profiles)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "The public Specialists Directory connects customers with verified clinics. "
        "Admins control who appears in the directory via the Manage Profiles console:"
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)
    
    add_bullet_point(doc, "Onboarding Pipeline: ", "When a partner applies, their profile is submitted to the database. Navigate to 'Manage Profiles' under 'Operations & CMS' in the sidebar.")
    add_bullet_point(doc, "Reviewing Profiles: ", "Click 'Edit' next to a specialist profile. Review their credentials, biography, and photos. Emojis and quotes are parsed automatically. Testimonial summaries are limited to 50 words to maintain directory page visual alignment.")
    add_bullet_point(doc, "Approval Status: ", "Toggle their approval checkmarks. Approved specialists show up immediately on the public B2C search grid. Verification checkmark badges are removed on directory cards to provide a cleaner, editorial style.")
    add_bullet_point(doc, "Location Formatting: ", "The database extracts clean Town/City headings automatically from address inputs (e.g. mapping '24 High St, Cardiff, CF10' to 'Cardiff') to keep directory cards uniform.")

    # Section 2
    add_heading_styled(doc, "2. Web Leads CRM & Pipeline (Manage Leads)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "All incoming client enquiries are routed into the central Leads Dashboard. "
        "To access, select 'Manage Leads' in the sidebar. The panel is split into four distinct inquiry types:"
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)
    
    add_bullet_point(doc, "Partner Enquiries: ", "Contains applications from external clinics wanting to host a TBN Hub.")
    add_bullet_point(doc, "Academy Registrations: ", "Tracks signups for TBN Clinical Academy courses.")
    add_bullet_point(doc, "Quiz Enquiries: ", "Detailed logs from the public B2C Health Goals Assessment quiz, showing client answers, goals, and referral codes.")
    add_bullet_point(doc, "Customer Booking Leads: ", "Tracks direct booking appointments submitted by clients on specialist pages.")

    add_heading_styled(doc, "CRM Deal Conversion and Deletion Actions:", 3)
    add_bullet_point(doc, "Pipeline Stage Tracking: ", "Draggable Kanban cards let you transition deals from 'Lead' -> 'Contacted' -> 'Meeting Scheduled' -> 'Active Partner'.")
    add_bullet_point(doc, "Notes & Updates: ", "Admins can click on deals to add meeting logs, change deal titles, or reassign them to specific regional specialist reps.")
    add_bullet_point(doc, "Lead Deletion (Spam Cleanup): ", "A secure Server Action allows admins to click the trash icon to purge spam leads. This immediately runs a service-role query to clear database records and syncs client metrics in real time.")
    add_bullet_point(doc, "CSV Export: ", "Click 'Export to CSV' on any of the leads sub-tabs to download log sheets for spreadsheets auditing.")

    # Section 3
    add_heading_styled(doc, "3. Site Traffic & Geolocation Analytics (Stats & Analytics)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "TBN runs a zero-configuration, self-hosted traffic tracker. Because it logs pageviews directly to your "
        "database, it bypasses ad-blockers and captures crucial visitor statistics before Google Analytics is active."
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)

    add_bullet_point(doc, "Traffic Metrics: ", "Dashboard metrics report Page Views, Unique Visitors, Average Session Duration, and Bounce Rates.")
    add_bullet_point(doc, "7-Day Trend Chart: ", "A custom vector line graph plots traffic logs over a rolling week.")
    add_bullet_point(doc, "Device & Acquisition: ", "View desktop/mobile ratios and acquisition splits (Direct, Organic Search, Social, Referrals).")
    add_bullet_point(doc, "Live Log Terminal: ", "Scrolls incoming pageviews in real time showing path, device type, campaign codes, and visitor's geolocated city/country (resolved via ipapi.co).")
    add_bullet_point(doc, "Google Analytics 4 (GA4): ", "Enter your measurement ID (e.g. G-XXXXX) in the configuration card to activate global Google tracking scripts.")

    add_callout(doc, "Tip: To test charts and table layouts, toggle the 'Sandbox Simulator' switch on the analytics page to stream randomized traffic data into your dashboard log.")

    # Section 4
    add_heading_styled(doc, "4. Affiliate Leaderboards & Campaign Builder (Affiliate Link Data)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "TBN drives growth through a practitioner referral network. The 'Affiliate Link Data' tab "
        "gives you visibility over referral campaigns:"
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)
    
    add_bullet_point(doc, "Campaign Link Builder: ", "Select any specialist in the network to generate custom referral links. Copies pre-tagged campaign slugs (e.g. ?ref=doctor_name) linking to the main B2C home page or their specialist directory profile.")
    add_bullet_point(doc, "Leaderboard Metrics: ", "Ranks specialists by referred registrations, completed client quizzes, and active signups.")
    add_bullet_point(doc, "Transaction Audits: ", "A real-time ledger auditing referred signups, showing dates, emails, category types, and attributing specialist codes.")

    # Section 5
    add_heading_styled(doc, "5. Content Management (Manage News & Categories)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "Admins can publish updates, news, and research to the B2C frontend. "
        "Navigate to 'Manage News' to draft articles, select publishing dates, upload thumbnails, "
        "and edit body content. Use 'Manage Categories' to create content tags (e.g. 'Women's Health', 'Cardiology') "
        "that help users search and filter articles."
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)

    # Section 6
    add_heading_styled(doc, "6. Cloud Storage & Library Uploads (Onboarding Resources)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "The partner resource explorer merges static PDF guides with files uploaded directly by admins. "
        "Under any training page (Start Here, Academy, Testing, Protocols, or Zinzino):"
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)
    
    add_bullet_point(doc, "Upload Console: ", "Admins automatically see the document uploader panel on dynamic directories. Supported files: PDFs, DOCX, XLSX, and MP4 videos. Enforces a client-size limit of 100MB per file.")
    add_bullet_point(doc, "Supabase Storage: ", "Files are stored in the secure 'partner-resources' bucket and public database records are created.")
    add_bullet_point(doc, "Resource Deletion: ", "Click the trash icon next to custom uploads to delete the database entry and purge the file from the storage bucket. Admins can also delete static defaults to hide outdated resources, tracked locally by browser cache exclusions.")

    # Save
    doc.save("TBN_Admin_Tutorial.docx")
    print("Admin_Tutorial.docx created successfully!")


def build_partner_tutorial():
    doc = Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title Page / Header
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("TESTBASED NUTRITION (TBN)\n")
    apply_text_formatting(run_title, font_name="Arial", size_pt=24, color_rgb=RGBColor(139, 26, 26), bold=True)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(36)
    run_sub = p_sub.add_run("Partner Clinic & Health Practitioner User Guide")
    apply_text_formatting(run_sub, font_name="Arial", size_pt=14, color_rgb=RGBColor(100, 100, 100), italic=True)

    add_heading_styled(doc, "Welcome to the TBN Partner Network", 1)
    p_intro = doc.add_paragraph()
    p_intro_run = p_intro.add_run(
        "Welcome, Partner! The TestBased Nutrition (TBN) portal is your personal dashboard "
        "for clinic growth, clinical training, resource libraries, and client CRM management. "
        "Our method is built on objective testing data: Test. Target. Transform. "
        "This guide walks you through using the portal to build your practice, manage client leads, "
        "complete academy credentials, and review referral commissions."
    )
    apply_text_formatting(p_intro_run, font_name="Arial", size_pt=10)

    # Section 1
    add_heading_styled(doc, "1. Portal Dashboard & Onboarding Checklist (Start Here)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "Upon logging in, you land on the 4x2 modules dashboard. At the top of your workspace, "
        "you will see your progress stepper tracking four stages: Learn, Launch, Grow, Lead."
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)
    
    add_bullet_point(doc, "Start Here: ", "Navigate to the first tile. The split-pane layout guides you through 6 core onboarding lessons. Check items off your checklist, and mark lessons as complete to update your dashboard status bar.")
    add_bullet_point(doc, "My Plans: ", "Track target timelines, set clinic launches, and customize next actions.")

    # Section 2
    add_heading_styled(doc, "2. TBN Academy & Testing Hub (Training Modules)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "The Academy is where you build clinical confidence. You must complete core training "
        "and testing certifications before offering TBN services to the public:"
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)
    
    add_bullet_point(doc, "Core Training: ", "Covering TBN foundations, consultation skills, results reviews, and compliance. Watch webinars, study slides, and complete the module quizzes. Quizzes require an 80% score to pass.")
    add_bullet_point(doc, "Testing Hub: ", "Find instructions, preparation guidelines, and shipping protocols for all in-clinic, finger-prick, and phlebotomy lab tests (such as Vitamin D, HbA1c, Ferritin, Cortisol, and FSH).")
    add_bullet_point(doc, "Zinzino Training: ", "Learn product lines (BalanceOil+, ZinoBiotic+), compensation routes, and Fast Start goals (securing 25 client subscriptions in your first 90 days).")

    # Section 3
    add_heading_styled(doc, "3. Clinical Protocols & Library (Pathways & Protocol Hub)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "The Protocols Library contains clinical data templates. Tap 'Protocol Hub' or 'Pathways Hub' in the sidebar:"
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)
    
    add_bullet_point(doc, "7 Health Categories: ", "Explore protocols for Women's Health, Men's Health, Children's Health, Skin Health, Pain & Inflammation, Neurodivergence, and Sports Performance.")
    add_bullet_point(doc, "50+ Sub-Pathways: ", "Includes step-by-step preparation, testing codes, results reviews, and supplement guidance for conditions like Eczema, Male Fertility, ADHD Support, or Chronic Pain.")
    add_bullet_point(doc, "One-Click Bulk Downloads: ", "Click 'Download Folder ZIP' to instantly package and download all files (PDF guides, templates, posters) within a module for offline use in your clinic.")
    add_bullet_point(doc, "Accessing Files: ", "Click 'View File' on any explorer link to instantly preview the PDF or file in your browser, or 'Download' to save it.")

    # Section 4
    add_heading_styled(doc, "4. Marketing Hub & AI Prompts (Grow Business)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "The Marketing Hub gives you a complete clinic promotion toolkit. No marketing experience needed:"
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)
    
    add_bullet_point(doc, "Client Acquisition Campaigns: ", "Download ready-made email templates, WhatsApp messages, and print-ready clinic posters for local awareness, workshop nights, and hub days.")
    add_bullet_point(doc, "Scripts & Conversation Guides: ", "Get discovery call templates, mid-treatment upgrade scripts, and retest reminder workflows to convert clients to long-term plans.")
    add_bullet_point(doc, "AI Content Creator Prompts: ", "Copy pre-tested ChatGPT prompts tailored to your specialization to draft custom social media posts, newsletters, and blog articles instantly.")

    # Section 5
    add_heading_styled(doc, "5. Directory Profile Management (My Profile)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "To customize your presence in the public Specialists Directory, navigate to 'My Profile'. "
        "Here you can edit your clinic location, certifications, services, and profile picture. "
        "Provide a short summary bio and expert quotes. To ensure page readability, quotes and testimonials "
        "are programmatically capped at 50 words on your public profile card."
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)

    # Section 6
    add_heading_styled(doc, "6. Client Leads Dashboard & CRM Deal Board (Manage Leads)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "When public users complete the Health Goals Quiz or submit bookings on your profile page, "
        "they are automatically logged in your Leads Dashboard under your partner code. "
        "Admins can also assign leads directly to your clinic."
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)
    
    add_bullet_point(doc, "CRM Deal Board: ", "A personalized Kanban board containing clients assigned to you. Drag client deal cards from 'Lead' -> 'Contacted' -> 'Meeting Scheduled' -> 'Active Patient' as you move them through consultations.")
    add_bullet_point(doc, "Client Records: ", "Click a deal to add notes from consultation calls, update dates, or review their selected pathway goals. Partners are restricted to editing their own deals only to ensure client privacy.")
    add_bullet_point(doc, "New Deals: ", "Manually add walk-in clients to your pipeline. New deals are automatically assigned to your practitioner profile.")

    # Section 7
    add_heading_styled(doc, "7. Custom Referral Links & Commission tracking", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "Referrals drive recurring income. Under 'My Referrals' or 'Edit Profile', "
        "you can manage and share your custom link structure:"
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)

    add_bullet_point(doc, "Custom Referral Code: ", "Set a clean, memorable affiliate code (e.g. 'doctor_emma'). Codes are checked for uniqueness in the database.")
    add_bullet_point(doc, "Professional Shortlinks: ", "Copy campaign links to send to clients or post on social media. Links direct users to the B2C home page (?ref=code) or directly to your clinic booking page (?ref=code).")
    add_bullet_point(doc, "Attribution tracking: ", "When users click your links, a tracker cookie caches your ID. Leads from contact forms, quiz completions, and bookings are attributed to your account automatically.")
    add_bullet_point(doc, "Subscription Income: ", "Track client product purchases and monthly retest cycles (120 days) directly from your Zinzino partner portal link.")

    # Section 8
    add_heading_styled(doc, "8. MLM Team Downline Tree (Team Downline)", 2)
    p = doc.add_paragraph()
    r = p.add_run(
        "As you grow your network, you can onboard and mentor other clinics. "
        "Select 'Team Downline' in the sidebar to view your organizational structure. "
        "The recursive visual tree maps your sponsored clinics, generation tiers, practitioner ranks, "
        "and active monthly enrollment volumes, making it easy to track regional performance commissions."
    )
    apply_text_formatting(r, font_name="Arial", size_pt=10)

    # Save
    doc.save("TBN_Partner_Tutorial.docx")
    print("Partner_Tutorial.docx created successfully!")


if __name__ == "__main__":
    build_admin_tutorial()
    build_partner_tutorial()
