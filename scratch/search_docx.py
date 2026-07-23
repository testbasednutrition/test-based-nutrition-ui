import docx
import os

files = [
    "TBN_Admin_Tutorial.docx",
    "TBN_Partner_Tutorial.docx",
    "TBN_README.docx",
    "TBN_Master_Structure (3).docx"
]

for f in files:
    if os.path.exists(f):
        doc = docx.Document(f)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    fullText.append(cell.text)
        text = "\n".join(fullText)
        print(f"--- {f} ---")
        for line in text.split("\n"):
            if "specialist" in line.lower() or "ambassador" in line.lower() or "form" in line.lower() or "profile" in line.lower():
                print(line)
