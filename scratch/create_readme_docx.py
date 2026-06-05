import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    # Set the background color of a cell
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def apply_text_formatting(run, font_name="Arial", size_pt=10, color_rgb=None, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_heading_styled(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    
    if level == 1:
        apply_text_formatting(run, font_name="Arial", size_pt=18, color_rgb=RGBColor(139, 26, 26), bold=True)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        apply_text_formatting(run, font_name="Arial", size_pt=14, color_rgb=RGBColor(139, 26, 26), bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        apply_text_formatting(run, font_name="Arial", size_pt=11.5, color_rgb=RGBColor(50, 50, 50), bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p

def add_bullet_point(doc, bold_prefix, text_body):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.4)
    
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        apply_text_formatting(run_bold, font_name="Arial", size_pt=10, bold=True)
    
    run_text = p.add_run(text_body)
    apply_text_formatting(run_text, font_name="Arial", size_pt=10)
    return p

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F5F5F5") # Light grey code background
    
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    run = p.add_run(code_text)
    apply_text_formatting(run, font_name="Courier New", size_pt=9, color_rgb=RGBColor(40, 40, 40))
    doc.add_paragraph().paragraph_format.space_after = Pt(2) # Margin below table

def build_readme_docx():
    doc = Document()
    
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Header / Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("TESTBASED NUTRITION (TBN)\n")
    apply_text_formatting(run_title, font_name="Arial", size_pt=24, color_rgb=RGBColor(139, 26, 26), bold=True)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Ecosystem Architecture & Master Documentation")
    apply_text_formatting(run_sub, font_name="Arial", size_pt=14, color_rgb=RGBColor(100, 100, 100), italic=True)

    # Divider line
    p_hr = doc.add_paragraph()
    p_hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_hr = p_hr.add_run("____________________________________________________")
    apply_text_formatting(run_hr, font_name="Arial", size_pt=10, color_rgb=RGBColor(200, 200, 200))
    p_hr.paragraph_format.space_after = Pt(20)

    # Intro Paragraph
    p_intro = doc.add_paragraph()
    p_intro_run = p_intro.add_run(
        "This document contains the complete structural overview of the TestBased Nutrition (TBN) ecosystem. "
        "This platform integrates the B2C Customer Portal with the B2B administrative back-office, automating "
        "practitioner onboarding, client leads tracking, CRM pipeline deals, and self-hosted traffic analytics."
    )
    apply_text_formatting(p_intro_run, font_name="Arial", size_pt=10)

    # 1. Architecture Flow
    add_heading_styled(doc, "Ecosystem Overview & Architecture", 1)
    
    # RENDER THE MERMAID FLOW GRAPH AS A BEAUTIFUL WORD TABLE
    p_flow_desc = doc.add_paragraph()
    p_flow_run = p_flow_desc.add_run("Below is the structural data flow between the public B2C app, the B2B Partner Portal, and the shared backend:")
    apply_text_formatting(p_flow_run, font_name="Arial", size_pt=10, italic=True)
    p_flow_desc.paragraph_format.space_after = Pt(8)

    # 4x2 styled table diagram
    table_diag = doc.add_table(rows=4, cols=2)
    table_diag.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Row 1: Header (merged)
    cell_h = table_diag.cell(0, 0)
    cell_h.merge(table_diag.cell(0, 1))
    set_cell_background(cell_h, "8B1A1A") # TBN Deep Red
    p_h = cell_h.paragraphs[0]
    p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_h = p_h.add_run("TBN ECOSYSTEM DATA FLOW")
    apply_text_formatting(run_h, font_name="Arial", size_pt=12, color_rgb=RGBColor(255, 255, 255), bold=True)
    
    # Row 2: Sub-headers
    cell_sh1 = table_diag.cell(1, 0)
    set_cell_background(cell_sh1, "F7F5F0")
    p_sh1 = cell_sh1.paragraphs[0]
    p_sh1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sh1 = p_sh1.add_run("Public B2C Frontend App")
    apply_text_formatting(run_sh1, font_name="Arial", size_pt=10, color_rgb=RGBColor(139, 26, 26), bold=True)

    cell_sh2 = table_diag.cell(1, 1)
    set_cell_background(cell_sh2, "F7F5F0")
    p_sh2 = cell_sh2.paragraphs[0]
    p_sh2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sh2 = p_sh2.add_run("B2B Partner Portal")
    apply_text_formatting(run_sh2, font_name="Arial", size_pt=10, color_rgb=RGBColor(139, 26, 26), bold=True)

    # Row 3: Items list
    cell_it1 = table_diag.cell(2, 0)
    p_it1 = cell_it1.paragraphs[0]
    p_it1.paragraph_format.left_indent = Inches(0.1)
    p_it1.paragraph_format.space_before = Pt(4)
    p_it1.paragraph_format.space_after = Pt(4)
    run_it1 = p_it1.add_run(
        "- Clinical Pathways (Sends pageviews)\n"
        "- Specialists Directory (Searches profiles)\n"
        "- B2C Health Goals Quiz (Saves quiz leads)\n"
        "- Route Tracker (Logs geolocation views)\n"
        "- Copy & DevTools Protection (Secures code)"
    )
    apply_text_formatting(run_it1, font_name="Arial", size_pt=9.5)

    cell_it2 = table_diag.cell(2, 1)
    p_it2 = cell_it2.paragraphs[0]
    p_it2.paragraph_format.left_indent = Inches(0.1)
    p_it2.paragraph_format.space_before = Pt(4)
    p_it2.paragraph_format.space_after = Pt(4)
    run_it2 = p_it2.add_run(
        "- Onboarding Checklists (Learn & Launch)\n"
        "- CRM Leads Dashboard (Manages deals Kanban)\n"
        "- MLM Downline Tree (Renders hierarchy lines)\n"
        "- Stats & Analytics Dashboard (Graphs views)\n"
        "- News CMS Engine (Publishes updates)"
    )
    apply_text_formatting(run_it2, font_name="Arial", size_pt=9.5)

    # Row 4: Backend Shared (merged)
    cell_b = table_diag.cell(3, 0)
    cell_b.merge(table_diag.cell(3, 1))
    set_cell_background(cell_b, "EFEFEF")
    p_b = cell_b.paragraphs[0]
    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_b = p_b.add_run("Supabase Shared Cloud Backend (PostgreSQL, Storage Buckets, Auth)")
    apply_text_formatting(run_b, font_name="Arial", size_pt=9.5, color_rgb=RGBColor(50, 50, 50), bold=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # B2C details
    add_heading_styled(doc, "1. The Public B2C Client App", 2)
    add_bullet_point(doc, "7 Health Pathways: ", "Deep-dive clinical content for Women's Health, Men's Health, Sports Performance, Skin Health, Pain & Fatigue, Neurodivergence, and Children's Health.")
    add_bullet_point(doc, "Directory & Search: ", "A 4-column specialist list grid with location search, automatic city extraction, and full text wrapping.")
    add_bullet_point(doc, "B2C Health Quiz: ", "A multi-step intake form collecting user goals and referral campaign slugs.")
    add_bullet_point(doc, "Route Analytics Tracker: ", "Logs visitor IDs, devices, referral campaign keys, and visitor city/country details.")

    # B2B details
    add_heading_styled(doc, "2. The B2B Partner Portal", 2)
    add_bullet_point(doc, "Onboarding & Academy: ", "checklist steps (Learn, Launch, Grow, Lead), webinars, lecture slides, and 80% passing quizzes.")
    add_bullet_point(doc, "CRM Leads Pipeline: ", "Draggable Kanban board to track leads from intake to active patient deals, with CSV spreadsheet download.")
    add_bullet_point(doc, "MLM Downline Tree: ", "A visual tree displaying team hierarchies, sponsor generation branches, monthly volumes, and ranks.")
    add_bullet_point(doc, "Analytics Dashboard: ", "Graphs 7-day traffic trends, channel sources, and live log terminal captures.")

    # Business Benefits
    add_heading_styled(doc, "Key Benefits & Business Logic", 1)
    
    add_heading_styled(doc, "1. For Clients & Patients (B2C Benefits)", 3)
    add_bullet_point(doc, "Objective Diagnostics: ", "Bypasses guesswork by utilizing point-of-care finger-prick and advanced lab testing.")
    add_bullet_point(doc, "Clear Timeline: ", "Unified 6-stage protocol bars (Discover, Test, Target, Transform, Retest, Escalate) provide path guidelines.")

    add_heading_styled(doc, "2. For Practitioners & Partners (B2B Benefits)", 3)
    add_bullet_point(doc, "Structured Business Model: ", "Step-by-step onboarding plan guides practitioners from setup to running regional hubs.")
    add_bullet_point(doc, "Compounding Income: ", "Builds monthly recurring revenue by enrolling clients in product subscriptions (BalanceOil+, ZinoBiotic+).")

    add_heading_styled(doc, "3. For Platform Administrators", 3)
    add_bullet_point(doc, "Self-Hosted Analytics: ", "Tracker logs pageviews directly to the database, ensuring zero dependencies on cookies before GA4 is set up.")
    add_bullet_point(doc, "Content Security: ", "DevTools blockades, click/selection interceptors, and frosted-glass alerts prevent data scraping.")

    # Testing Tier Table
    add_heading_styled(doc, "Testing & Screening Framework", 1)
    
    table_test = doc.add_table(rows=4, cols=4)
    table_test.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Headers
    headers = ["Testing Tier", "Location", "Key Bio-Markers", "Clinical Action"]
    for col_idx, text in enumerate(headers):
        cell = table_test.cell(0, col_idx)
        set_cell_background(cell, "8B1A1A")
        p = cell.paragraphs[0]
        run = p.add_run(text)
        apply_text_formatting(run, font_name="Arial", size_pt=9.5, color_rgb=RGBColor(255, 255, 255), bold=True)

    # Data Rows
    row_data = [
        ("Foundational Testing", "In-Clinic or Online", "Omega Balance Test, Gut Microbiome Test, FSH", "Establish baseline cellular inflammation, fatty acid profile, hormone signalling, and microbiome diversity."),
        ("Baseline Screening", "Finger-Prick POC", "Vitamin D, HbA1c, Ferritin, CRP, Cortisol, Cystatin C, Folate, HCG-b, AMH, Progesterone, NT-proBNP, RSV, RF", "Rapid, 15-minute diagnostic checks performed during the initial physical consultation."),
        ("Advanced Screening", "Phlebotomy Lab", "Testosterone, Vitamin B12, FSH, Thyroid (TSH)", "In-depth biochemical analysis triggered when symptoms suggest systemic imbalance.")
    ]

    for row_idx, data in enumerate(row_data):
        for col_idx, text in enumerate(data):
            cell = table_test.cell(row_idx + 1, col_idx)
            if row_idx % 2 == 0:
                set_cell_background(cell, "FFFFFF")
            else:
                set_cell_background(cell, "F9F9F9") # Alternating row colors
            p = cell.paragraphs[0]
            run = p.add_run(text)
            apply_text_formatting(run, font_name="Arial", size_pt=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Tech Stack
    add_heading_styled(doc, "Tech Stack", 1)
    add_bullet_point(doc, "Frontend: ", "React / Vite (Main Site), Next.js App Router (Partner Hub)")
    add_bullet_point(doc, "Database: ", "Supabase (PostgreSQL, Storage, RLS policies)")
    add_bullet_point(doc, "Styling: ", "Tailwind CSS & custom CSS palette (Sand, cream, and deep-red #8B1A1A)")
    add_bullet_point(doc, "Components: ", "shadcn/ui & Radix UI primitives")

    # Getting Started
    add_heading_styled(doc, "Getting Started (Development)", 1)
    
    add_heading_styled(doc, "1. Start B2C Website", 3)
    add_code_block(doc, "cd test-based-nutrition-ui\nnpm install\nnpm run dev")

    add_heading_styled(doc, "2. Start Partner Portal", 3)
    add_code_block(doc, "cd partner-hub\npnpm install\npnpm dev")

    # Database Schema
    add_heading_styled(doc, "Database Schema Setup", 1)
    p_schema = doc.add_paragraph()
    r_schema = p_schema.add_run("Below is the SQL table schema for pageview tracking logs:")
    apply_text_formatting(r_schema, font_name="Arial", size_pt=10, italic=True)
    
    sql_text = (
        "CREATE TABLE public.page_views (\n"
        "    id uuid DEFAULT gen_random_uuid() NOT NULL PRIMARY KEY,\n"
        "    created_at timestamp with time zone DEFAULT timezone('utc'::text, now()) NOT NULL,\n"
        "    path text NOT NULL,\n"
        "    referrer text,\n"
        "    visitor_id text NOT NULL,\n"
        "    device_type text,\n"
        "    city text,\n"
        "    country text,\n"
        "    campaign_code text\n"
        ");"
    )
    add_code_block(doc, sql_text)

    # Save
    doc.save("TBN_README.docx")
    print("TBN_README.docx created successfully!")

if __name__ == "__main__":
    build_readme_docx()
